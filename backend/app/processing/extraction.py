from __future__ import annotations

import json
import logging
import mimetypes
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Iterable

from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

MarkdownHeading = re.compile(r"^(#+)\s+(?P<title>.+)$")


class ExtractionError(Exception):
    """Raised when a document cannot be extracted or parsed."""


@dataclass
class ExtractedSection:
    index: int
    title: str | None
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "index": self.index,
            "title": self.title,
            "content": self.content,
            "metadata": self.metadata,
        }


@dataclass
class ExtractedDocument:
    document_path: str
    source_extension: str
    content_type: str
    sections: list[ExtractedSection]
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "document_path": self.document_path,
            "source_extension": self.source_extension,
            "content_type": self.content_type,
            "section_count": len(self.sections),
            "metadata": self.metadata,
            "sections": [section.to_dict() for section in self.sections],
        }

    def to_json(self, *, indent: int | None = None) -> str:
        return json.dumps(self.to_dict(), indent=indent, ensure_ascii=False)


class DocumentExtractor:
    """High-level text extraction orchestrator for supported document formats."""

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".md",
        ".markdown",
        ".txt",
        ".html",
        ".htm",
        ".json",
        ".xml",
    }
    OCR_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".tif"}

    def __init__(
        self,
        *,
        use_ocr: bool = False,
        ocr_lang: str = "eng",
        min_section_length: int = 20,
    ):
        self.use_ocr = use_ocr
        self.ocr_lang = ocr_lang
        self.min_section_length = min_section_length

    def extract(self, path: str | Path) -> ExtractedDocument:
        document_path = Path(path)
        if not document_path.exists():
            raise ExtractionError(f"Document not found: {document_path}")
        if not document_path.is_file():
            raise ExtractionError(f"Path must be a file: {document_path}")

        extension = document_path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS and (
            extension not in self.OCR_EXTENSIONS
        ):
            allowed = ", ".join(sorted(self.SUPPORTED_EXTENSIONS | self.OCR_EXTENSIONS))
            raise ExtractionError(
                f"Unsupported extension '{extension}'. Allowed: {allowed}"
            )

        extractor = self._resolve_extractor(extension)
        sections = extractor(document_path)

        if not sections:
            raise ExtractionError(
                f"No textual content could be extracted from {document_path.name}."
            )

        content_type = (
            mimetypes.guess_type(document_path.name)[0] or "application/octet-stream"
        )
        metadata = {
            "extracted_at": datetime.utcnow().isoformat() + "Z",
            "use_ocr": self.use_ocr,
            "source_name": document_path.name,
        }

        return ExtractedDocument(
            document_path=str(document_path),
            source_extension=extension,
            content_type=content_type,
            sections=sections,
            metadata=metadata,
        )

    # --------------------------------------------------------------------- #
    # Individual format extractors
    # --------------------------------------------------------------------- #
    def _resolve_extractor(
        self, extension: str
    ) -> Callable[[Path], list[ExtractedSection]]:
        if extension in {".pdf"}:
            return self._extract_pdf
        if extension in {".docx"}:
            return self._extract_docx
        if extension in {".md", ".markdown"}:
            return self._extract_markdown
        if extension in {".txt", ".json"}:
            return self._extract_plain_text
        if extension in {".html", ".htm"}:
            return self._extract_html
        if extension in {".xml"}:
            return self._extract_xml
        if extension in self.OCR_EXTENSIONS:
            if not self.use_ocr:
                raise ExtractionError(
                    f"OCR is disabled; enable --ocr to process {extension} files."
                )
            return self._extract_image_ocr
        raise ExtractionError(f"No extractor available for extension {extension}")

    def _extract_pdf(self, path: Path) -> list[ExtractedSection]:
        from PyPDF2 import PdfReader  # local import to reduce startup time

        sections: list[ExtractedSection] = []
        reader = PdfReader(str(path))
        for index, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = self._normalize_whitespace(text)
            if not text and self.use_ocr:
                logger.debug("PDF page %s empty, attempting OCR fallback.", index + 1)
                text = self._extract_pdf_page_with_ocr(path, index)
            if not text:
                continue
            sections.append(
                ExtractedSection(
                    index=index,
                    title=f"Page {index + 1}",
                    content=text,
                    metadata={"source": "pdf", "page_number": index + 1},
                )
            )
        return sections

    def _extract_docx(self, path: Path) -> list[ExtractedSection]:
        from docx import Document as DocxDocument

        document = DocxDocument(str(path))
        sections: list[ExtractedSection] = []
        buffer: list[str] = []
        current_title: str | None = None
        section_index = 0

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = getattr(paragraph.style, "name", "") or ""
            if style_name.startswith("Heading"):
                if buffer:
                    sections.append(
                        ExtractedSection(
                            index=section_index,
                            title=current_title or f"Section {section_index + 1}",
                            content=self._join_buffer(buffer),
                            metadata={
                                "source": "docx",
                                "style": style_name,
                            },
                        )
                    )
                    buffer.clear()
                    section_index += 1

                current_title = text
                continue

            buffer.append(text)

        if buffer:
            sections.append(
                ExtractedSection(
                    index=section_index,
                    title=current_title or f"Section {section_index + 1}",
                    content=self._join_buffer(buffer),
                    metadata={"source": "docx"},
                )
            )

        return sections

    def _extract_markdown(self, path: Path) -> list[ExtractedSection]:
        text = path.read_text(encoding="utf-8")
        lines = text.splitlines()
        sections: list[ExtractedSection] = []
        buffer: list[str] = []
        current_title: str | None = None
        for raw_line in lines:
            line = raw_line.rstrip()
            heading_match = MarkdownHeading.match(line)
            if heading_match:
                self._flush_section(
                    sections,
                    buffer,
                    len(sections),
                    current_title,
                    source="markdown",
                )
                current_title = heading_match.group("title").strip()
                buffer.clear()
                continue

            buffer.append(line)

        self._flush_section(
            sections,
            buffer,
            len(sections),
            current_title,
            source="markdown",
        )

        if not sections:
            content = self._normalize_whitespace(text)
            sections.append(
                ExtractedSection(
                    index=0,
                    title=None,
                    content=content,
                    metadata={"source": "markdown"},
                )
            )
        return sections

    def _extract_html(self, path: Path) -> list[ExtractedSection]:
        html_text = path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html_text, "html.parser")
        headings = soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])

        if not headings:
            body_text = soup.get_text(" ", strip=True)
            return [
                ExtractedSection(
                    index=0,
                    title="Document Body",
                    content=self._normalize_whitespace(body_text),
                    metadata={"source": "html"},
                )
            ]

        sections: list[ExtractedSection] = []
        for idx, heading in enumerate(headings):
            title = heading.get_text(" ", strip=True)
            content_parts: list[str] = []
            for sibling in heading.next_siblings:
                if (
                    getattr(sibling, "name", "") in {"h1", "h2", "h3", "h4", "h5", "h6"}
                ):
                    break
                if getattr(sibling, "name", None) in {None, "p", "div", "span", "li"}:
                    text = (
                        sibling.get_text(" ", strip=True)
                        if hasattr(sibling, "get_text")
                        else str(sibling).strip()
                    )
                    if text:
                        content_parts.append(text)
            sections.append(
                ExtractedSection(
                    index=idx,
                    title=title,
                    content=self._normalize_whitespace("\n".join(content_parts)),
                    metadata={"source": "html", "heading": heading.name},
                )
            )

        return sections

    def _extract_xml(self, path: Path) -> list[ExtractedSection]:
        """Extract text from XML files, preserving structure where possible."""
        # Try to detect encoding from XML declaration or file content
        try:
            # Read as bytes first to check for BOM
            raw_bytes = path.read_bytes()
            
            # Check for UTF-8 BOM
            if raw_bytes[:3] == b'\xef\xbb\xbf':
                xml_text = raw_bytes[3:].decode('utf-8')
            # Check for UTF-16 BOM
            elif raw_bytes[:2] == b'\xff\xfe':
                xml_text = raw_bytes[2:].decode('utf-16-le')
            elif raw_bytes[:2] == b'\xfe\xff':
                xml_text = raw_bytes[2:].decode('utf-16-be')
            else:
                # Try to read as UTF-8, fallback to latin-1 if that fails
                try:
                    xml_text = raw_bytes.decode('utf-8')
                except UnicodeDecodeError:
                    # Try to detect encoding from XML declaration
                    xml_decl = raw_bytes[:200].decode('latin-1', errors='ignore')
                    if 'encoding=' in xml_decl:
                        import re
                        match = re.search(r'encoding=["\']([^"\']+)["\']', xml_decl)
                        if match:
                            encoding = match.group(1).lower()
                            try:
                                xml_text = raw_bytes.decode(encoding)
                            except (UnicodeDecodeError, LookupError):
                                xml_text = raw_bytes.decode('utf-8', errors='replace')
                        else:
                            xml_text = raw_bytes.decode('utf-8', errors='replace')
                    else:
                        xml_text = raw_bytes.decode('utf-8', errors='replace')
        except Exception:
            # Fallback to simple UTF-8 read
            xml_text = path.read_text(encoding="utf-8", errors='replace')
        
        soup = BeautifulSoup(xml_text, "xml")  # Use XML parser
        
        # Check if this is a Microsoft Office XML package (Word.Document)
        # These have a structure like: <pkg:package><pkg:part pkg:name="/word/document.xml">...
        is_office_xml = (
            soup.find("pkg:package") is not None or 
            soup.find("package") is not None or
            'mso-application' in xml_text.lower() or
            'word.document' in xml_text.lower()
        )
        
        if is_office_xml:
            # This is a Microsoft Office XML package - extract from document part
            return self._extract_office_xml(soup, xml_text)
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Try to find meaningful structure (sections, articles, parts, etc.)
        # Look for common regulation/document structure elements
        structure_tags = ["section", "article", "part", "chapter", "title", "division", "regulation"]
        sections: list[ExtractedSection] = []
        section_index = 0
        
        # First, try to find structured elements
        for tag_name in structure_tags:
            elements = soup.find_all(tag_name, recursive=True)
            if elements:
                for elem in elements:
                    title = elem.get("title") or elem.get("name") or elem.get("id") or tag_name.title()
                    content = elem.get_text(" ", strip=True)
                    if content and len(content) >= self.min_section_length:
                        sections.append(
                            ExtractedSection(
                                index=section_index,
                                title=title,
                                content=self._normalize_whitespace(content),
                                metadata={"source": "xml", "tag": tag_name},
                            )
                        )
                        section_index += 1
                break  # Use first matching structure
        
        # If no structured elements found, extract by paragraphs or divs
        if not sections:
            paragraphs = soup.find_all(["p", "div", "para", "text"])
            if paragraphs:
                for idx, para in enumerate(paragraphs):
                    text = para.get_text(" ", strip=True)
                    if text and len(text) >= self.min_section_length:
                        title = para.get("title") or para.get("id") or f"Section {idx + 1}"
                        sections.append(
                            ExtractedSection(
                                index=idx,
                                title=title,
                                content=self._normalize_whitespace(text),
                                metadata={"source": "xml"},
                            )
                        )
        
        # Fallback: extract all text as single section
        if not sections:
            all_text = soup.get_text(" ", strip=True)
            if all_text:
                sections.append(
                    ExtractedSection(
                        index=0,
                        title="Document Content",
                        content=self._normalize_whitespace(all_text),
                        metadata={"source": "xml"},
                    )
                )
        
        return sections
    
    def _extract_office_xml(self, soup: BeautifulSoup, xml_text: str) -> list[ExtractedSection]:
        """Extract text from Microsoft Office XML package (Word.Document format)."""
        sections: list[ExtractedSection] = []
        
        # Find the main document part (/word/document.xml)
        # Look for pkg:part with name="/word/document.xml" or similar
        document_parts = []
        
        # Try different namespace prefixes
        for part in soup.find_all(["pkg:part", "part"]):
            part_name = part.get("pkg:name") or part.get("name", "")
            if "/word/document.xml" in part_name or "/word/document" in part_name:
                document_parts.append(part)
        
        if not document_parts:
            # Fallback: look for any part containing Word document XML
            for part in soup.find_all(["pkg:part", "part"]):
                part_xml = part.find("pkg:xmlData") or part
                if part_xml.find("w:document") or part_xml.find("document"):
                    document_parts.append(part_xml)
        
        # Extract text from document parts
        for part_idx, part in enumerate(document_parts):
            # Get the XML data within the part
            xml_data = part.find("pkg:xmlData") or part
            
            # Find all text nodes in Word XML (w:t elements contain text)
            # Word uses namespace w: for wordprocessingml
            text_elements = xml_data.find_all(["w:t", "t"])
            
            if text_elements:
                # Group text by paragraphs
                paragraphs = xml_data.find_all(["w:p", "p"])
                
                for para_idx, para in enumerate(paragraphs):
                    # Extract text from this paragraph
                    para_text = para.get_text(" ", strip=True)
                    if para_text and len(para_text) >= self.min_section_length:
                        # Try to get heading style or title
                        heading = para.find(["w:pStyle", "pStyle"])
                        title = "Paragraph"
                        if heading:
                            style_val = heading.get("w:val") or heading.get("val", "")
                            if style_val:
                                title = f"Section {style_val}"
                        
                        sections.append(
                            ExtractedSection(
                                index=len(sections),
                                title=title,
                                content=self._normalize_whitespace(para_text),
                                metadata={"source": "xml", "format": "office_xml", "part": part_idx},
                            )
                        )
                
                # If no paragraphs found, extract all text elements
                if not sections and text_elements:
                    all_text = " ".join([elem.get_text(" ", strip=True) for elem in text_elements if elem.get_text(strip=True)])
                    if all_text and len(all_text) >= self.min_section_length:
                        sections.append(
                            ExtractedSection(
                                index=0,
                                title="Document Content",
                                content=self._normalize_whitespace(all_text),
                                metadata={"source": "xml", "format": "office_xml"},
                            )
                        )
        
        # If still no sections, fallback to extracting all text from the XML
        if not sections:
            # Remove package structure and extract meaningful text
            all_text = soup.get_text(" ", strip=True)
            # Filter out obvious metadata
            lines = [line.strip() for line in all_text.split("\n") if line.strip()]
            meaningful_lines = [
                line for line in lines 
                if not line.startswith("Document ID") 
                and not line.startswith("DocumentLibrary")
                and len(line) > 20
            ]
            
            if meaningful_lines:
                content = " ".join(meaningful_lines)
                if len(content) >= self.min_section_length:
                    sections.append(
                        ExtractedSection(
                            index=0,
                            title="Document Content",
                            content=self._normalize_whitespace(content),
                            metadata={"source": "xml", "format": "office_xml"},
                        )
                    )
        
        return sections

    def _extract_plain_text(self, path: Path) -> list[ExtractedSection]:
        text = path.read_text(encoding="utf-8")
        content = self._normalize_whitespace(text)
        if not content:
            return []
        return [
            ExtractedSection(
                index=0,
                title="Document Body",
                content=content,
                metadata={"source": "text"},
            )
        ]

    def _extract_image_ocr(self, path: Path) -> list[ExtractedSection]:
        if not self.use_ocr:
            return []
        try:
            from PIL import Image
            import pytesseract
        except ImportError as exc:  # pragma: no cover - guarded by dependency
            raise ExtractionError(
                "pytesseract and pillow are required for OCR extraction."
            ) from exc

        with Image.open(path) as image:
            text = pytesseract.image_to_string(image, lang=self.ocr_lang)

        content = self._normalize_whitespace(text)
        if not content:
            return []

        return [
            ExtractedSection(
                index=0,
                title="OCR Content",
                content=content,
                metadata={"source": "ocr", "lang": self.ocr_lang},
            )
        ]

    def _extract_pdf_page_with_ocr(self, path: Path, page_index: int) -> str:
        try:
            import fitz  # type: ignore
        except ImportError:  # pragma: no cover - optional dependency
            logger.warning(
                "PyMuPDF (fitz) not installed; skipping OCR fallback for %s page %s",
                path.name,
                page_index + 1,
            )
            return ""

        try:
            import pytesseract
        except ImportError:  # pragma: no cover - optional dependency
            logger.warning("pytesseract not installed; skipping OCR fallback.")
            return ""

        doc = fitz.open(path)  # pragma: no cover - requires optional dependency
        page = doc.load_page(page_index)
        pix = page.get_pixmap()
        from PIL import Image  # pragma: no cover
        from io import BytesIO  # pragma: no cover

        image = Image.open(BytesIO(pix.tobytes()))
        text = pytesseract.image_to_string(image, lang=self.ocr_lang)
        return self._normalize_whitespace(text)

    # ------------------------------------------------------------------ #
    # Helpers
    # ------------------------------------------------------------------ #
    def _flush_section(
        self,
        sections: list[ExtractedSection],
        buffer: list[str],
        index: int,
        title: str | None,
        *,
        source: str,
    ) -> None:
        content = self._normalize_whitespace(self._join_buffer(buffer))
        if not content or len(content) < self.min_section_length:
            buffer.clear()
            return
        sections.append(
            ExtractedSection(
                index=index,
                title=title,
                content=content,
                metadata={"source": source},
            )
        )
        buffer.clear()

    @staticmethod
    def _join_buffer(buffer: Iterable[str]) -> str:
        return "\n".join(line.strip() for line in buffer if line.strip())

    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        normalized = re.sub(r"[ \t]+", " ", text)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()


